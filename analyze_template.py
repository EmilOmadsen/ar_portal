"""
Analyze Word template structure to identify all placeholders and their context
"""
from docx import Document
import re

doc = Document('static/contract_templates/50_50 template med placeholders.docx')

print("=" * 80)
print("TEMPLATE STRUCTURE ANALYSIS")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    # Check for Jinja2 placeholders
    if '{{' in text or '{%' in text:
        print(f"\n[Para {i}]")
        print(f"Text: {text[:150]}")
        
        # Extract placeholders
        placeholders = re.findall(r'\{\{.*?\}\}', text)
        conditionals = re.findall(r'\{%.*?%\}', text)
        
        if placeholders:
            print(f"Placeholders: {placeholders}")
        if conditionals:
            print(f"Conditionals: {conditionals}")

print("\n" + "=" * 80)
print("SUMMARY")
print("=" * 80)

all_text = '\n'.join([p.text for p in doc.paragraphs])

# Count structures
for_loops = len(re.findall(r'\{% for ', all_text))
endfors = len(re.findall(r'\{% endfor %\}', all_text))
ifs = len(re.findall(r'\{% if ', all_text))
endifs = len(re.findall(r'\{% endif %\}', all_text))

print(f"FOR loops: {for_loops}")
print(f"ENDFOR: {endfors}")
print(f"IF statements: {ifs}")
print(f"ENDIF: {endifs}")

# Find all unique placeholders
all_placeholders = set(re.findall(r'\{\{(.*?)\}\}', all_text))
print(f"\nUnique placeholders ({len(all_placeholders)}):")
for p in sorted(all_placeholders):
    print(f"  - {{{{{p}}}}}")
