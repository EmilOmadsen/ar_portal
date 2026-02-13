"""
Add Jinja2 placeholders to the contract template
"""
from docx import Document
from pathlib import Path
import re

def add_placeholders_to_contract():
    # Load the original contract
    input_file = Path("static/contract_templates/50 - ENG - REC_PUB - SUNDAY.docx")
    output_file = Path("static/contract_templates/50_50 template med placeholders.docx")
    
    print(f"Loading contract from: {input_file}")
    doc = Document(input_file)
    
    print(f"Contract has {len(doc.paragraphs)} paragraphs")
    print("\nAdding placeholders...")
    
    replacements = [
        # Artist information
        ("Artist Name", "{{artist.stage_name}}"),
        ("Full Name", "{{artist.full_name}}"),
        ("Street Name and Number", "{{artist.address}}"),
        ("Zip Code/Postal Code/Fiscal Code", "{{artist.postcode}}"),
        ("City\nCountry", "{{artist.city}}\n{{artist.country}}"),
        ("Social Security ID (or other personal identification number from your government)", "{{artist.id_number}}"),
        ("E-mail address", "{{artist.email}}"),
        ("IPI Number (leave blank if you do not have it)", "{{artist.ipi}}"),
        
        # Project information
        ("Project Number:", "Project Code: {{project_code}}"),
        
        # Financial terms
        ("50%", "{{royalty_percent}}%"),
        ("fifty percent", "{{royalty_percent}} percent"),
    ]
    
    changes_made = 0
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        original_text = para.text
        new_text = original_text
        
        for old, new in replacements:
            if old in new_text:
                new_text = new_text.replace(old, new)
                if new_text != original_text:
                    changes_made += 1
                    print(f"  Replaced: '{old}' -> '{new}'")
        
        if new_text != original_text:
            # Clear existing runs and add new text
            para.clear()
            para.add_run(new_text)
    
    # Find where to add FOR loop for artists
    # Look for the artist section
    for i, para in enumerate(doc.paragraphs):
        if "Upon signature the following parties:" in para.text:
            print(f"\n  Adding artist loop at paragraph {i+9}")
            # Insert FOR loop before artist info
            doc.paragraphs[i+9].insert_paragraph_before("{% for artist in artists %}")
            break
    
    # Find where to end FOR loop (after artist info, before "WITNESSETH:")
    for i, para in enumerate(doc.paragraphs):
        if "WITNESSETH:" in para.text:
            print(f"  Closing artist loop before paragraph {i}")
            # Add the comma separator inside the loop, then close the loop
            doc.paragraphs[i].insert_paragraph_before("{% endfor %}")
            break
    
    # Note: Recordings section would need manual editing in LibreOffice
    # Skip automatic recording loop for now
    
    # Add conditional for publishing
    for i, para in enumerate(doc.paragraphs):
        if "WRITER" in para.text and "PUBLISHER" in para.text:
            print(f"\n  Adding publishing conditional at paragraph {i}")
            doc.paragraphs[i].insert_paragraph_before("{% if publishing_included %}")
            # Find the end of publishing section
            for j in range(i+1, min(i+50, len(doc.paragraphs))):
                if "Recording" in doc.paragraphs[j].text or "WITNESSETH" in doc.paragraphs[j].text:
                    print(f"  Closing publishing conditional before paragraph {j}")
                    doc.paragraphs[j].insert_paragraph_before("{% endif %}")
                    break
            break
    
    # Add marketing recoupment conditional with numeric + words format
    for i, para in enumerate(doc.paragraphs):
        if "marketing expenses" in para.text.lower() or "recoupable" in para.text.lower():
            print(f"\n  Adding marketing recoupment conditional around paragraph {i}")
            doc.paragraphs[i].insert_paragraph_before("{% if marketing_recoupment_enabled %}")
            # Clear original and replace with new format
            para.clear()
            para.add_run("Company shall be allowed to offset a recoupable cost of USD {{marketing_recoupment_amount_numeric}} ({{marketing_recoupment_amount_words}} American Dollars)")
            # Add option if enabled
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i+1].insert_paragraph_before("{% if marketing_option_enabled %} with an option of USD {{marketing_option_amount_numeric}} ({{marketing_option_amount_words}} American Dollars){% endif %}, all from royalties due to Artist for each Recording released under this agreement to compensate Company for marketing expenses incurred by Company. For the avoidance of doubt, this shall not apply to remixes or alternate versions.")
                doc.paragraphs[i+2].insert_paragraph_before("{% endif %}")
            break
    
    # Add standard advance section
    for i, para in enumerate(doc.paragraphs):
        if "advance" in para.text.lower() and "milestone" not in para.text.lower() and i > 50:  # Skip early mentions
            print(f"\n  Adding standard advance conditional at paragraph {i}")
            doc.paragraphs[i].insert_paragraph_before("{% if advance_enabled %}")
            para.clear()
            para.add_run("Company agrees to pay Artist an advance of USD {{advance_amount_numeric}} ({{advance_amount_words}} American Dollars) upon execution of this agreement.")
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i+1].insert_paragraph_before("{% endif %}")
            break
    
    # Add milestone advance section
    print(f"\n  Adding milestone advance section after standard advance")
    # Find a good spot to insert (after advances section)
    for i, para in enumerate(doc.paragraphs):
        if "advance" in para.text.lower() and i > 70:  # After standard advance
            doc.paragraphs[i+3].insert_paragraph_before("{% if milestone_enabled %}")
            doc.paragraphs[i+4].insert_paragraph_before("Milestone Advance: If the Recording achieves {{milestone_daily_streams}} daily streams for {{milestone_period_days}} consecutive days, Company agrees to pay an additional advance of USD {{milestone_advance_amount_numeric}} ({{milestone_advance_amount_words}} American Dollars).")
            doc.paragraphs[i+5].insert_paragraph_before("{% endif %}")
            break
    
    print(f"\n\nTotal changes made: {changes_made}")
    
    # Save the modified document
    doc.save(output_file)
    print(f"\nSaved template with placeholders to: {output_file}")
    
    # Verify the output
    verify_doc = Document(output_file)
    print(f"\nVerification:")
    print(f"  Output file has {len(verify_doc.paragraphs)} paragraphs")
    
    # Check for placeholders
    placeholder_count = 0
    for para in verify_doc.paragraphs:
        if "{{" in para.text or "{%" in para.text:
            placeholder_count += 1
    
    print(f"  Found {placeholder_count} paragraphs with Jinja2 placeholders")
    
    if placeholder_count > 0:
        print("\n  SUCCESS! Placeholders added successfully!")
        return True
    else:
        print("\n  WARNING: No placeholders found in output")
        return False

if __name__ == "__main__":
    try:
        success = add_placeholders_to_contract()
        if success:
            print("\nYou can now use this template in your contract generator!")
            print("Next step: Run python fix_template.py to fix any tag issues")
    except Exception as e:
        print(f"\nERROR: {e}")
        import traceback
        traceback.print_exc()
