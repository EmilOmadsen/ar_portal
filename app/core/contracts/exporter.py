from tempfile import NamedTemporaryFile
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def render_contract(template, context: dict):
    """
    Render the Word (.docx) template with Jinja context
    and return the file path of the generated document.
    
    Note: Images, headers, and footers in the template are automatically preserved.
    The docxtpl library maintains all formatting and embedded media.
    """
    # Render the template with the provided context
    # This replaces all {{placeholder}} tags while preserving images and formatting
    template.render(context, autoescape=True)

    # Save to temporary file
    tmp = NamedTemporaryFile(delete=False, suffix=".docx")
    template.save(tmp.name)

    # Yellow highlighting disabled
    # _add_highlighting_to_fields(tmp.name, context)

    return tmp.name


def _add_highlighting_to_fields(docx_path: str, context: dict):
    """
    Add yellow highlighting to runs in the document that match populated field values
    """
    try:
        doc = Document(docx_path)
        
        # Collect all populated values to highlight
        values_to_highlight = _extract_highlight_values(context)
        
        # Iterate through all paragraphs and runs
        for paragraph in doc.paragraphs:
            _highlight_values_in_paragraph(paragraph, values_to_highlight)
        
        # Also check in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _highlight_values_in_paragraph(paragraph, values_to_highlight)
        
        # Save the document back
        doc.save(docx_path)
    except Exception as e:
        print(f"Warning: Could not add highlighting to DOCX: {e}")
        # Don't fail contract generation if highlighting fails


def _highlight_values_in_paragraph(paragraph, values_to_highlight):
    """
    Highlight values within a paragraph by splitting and rebuilding runs
    """
    for value in values_to_highlight:
        if not value or len(value) < 2:
            continue
        
        # Convert paragraph text to lowercase for matching
        para_text_lower = paragraph.text.lower()
        value_lower = value.lower()
        
        # Check if value exists in paragraph
        if value_lower not in para_text_lower:
            continue
        
        # Find all occurrences and highlight them
        text = paragraph.text
        text_lower = text.lower()
        start_idx = 0
        
        while True:
            # Find next occurrence
            match_idx = text_lower.find(value_lower, start_idx)
            if match_idx == -1:
                break
            
            # Try to highlight this occurrence
            _highlight_run_in_range(paragraph, match_idx, len(value))
            start_idx = match_idx + len(value)


def _highlight_run_in_range(paragraph, start_char, length):
    """
    Find and highlight a run (or spans across runs) at a specific character position
    """
    current_pos = 0
    for run in paragraph.runs:
        run_len = len(run.text)
        run_end = current_pos + run_len
        
        # Check if this run contains part of the range
        if run_end > start_char and current_pos < start_char + length:
            _add_yellow_highlight(run)
        
        current_pos = run_end


def _extract_highlight_values(context: dict) -> list:
    """
    Extract all populated values from context that should be highlighted.
    This recursively gets all string and number values from the context.
    """
    values = []
    
    def extract_values(obj, visited=None):
        if visited is None:
            visited = set()
        
        # Avoid infinite recursion
        obj_id = id(obj)
        if obj_id in visited:
            return
        visited.add(obj_id)
        
        if isinstance(obj, dict):
            for key, value in obj.items():
                # Skip certain keys that aren't user-entered values
                if key.lower() in ['_', '__', 'type', 'class']:
                    continue
                extract_values(value, visited)
        elif isinstance(obj, (list, tuple)):
            for item in obj:
                extract_values(item, visited)
        elif isinstance(obj, (str, int, float)):
            if obj and str(obj).strip():
                values.append(str(obj))
        elif isinstance(obj, bool):
            # Skip boolean values as they're not content to highlight
            pass
    
    extract_values(context)
    
    # Remove duplicates and empty values, sort by length (longest first for better matching)
    values = list(set(values))
    values = [v for v in values if v.strip() and len(v.strip()) > 1]  # Skip single characters
    values.sort(key=len, reverse=True)
    
    return values


def _add_yellow_highlight(run):
    """
    Add yellow highlighting to a run using XML
    """
    try:
        # Create highlight element (yellow = 7 in Word XML)
        shd = parse_xml(r'<w:highlight {} w:val="yellow"/>'.format(nsdecls('w')))
        run._element.get_or_add_rPr().append(shd)
    except Exception as e:
        print(f"Could not add highlight to run: {e}")
